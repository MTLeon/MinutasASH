from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_text(cell, text: str, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _add_heading(document: DocxDocument, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)


def _add_empty_message(document: DocxDocument, message: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(message)
    run.italic = True


def generate_docx(
    minute: Any,
    metadata: dict,
    output_path: str | Path,
    company_name: str,
) -> Path:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MINUTA DE REUNIÓN")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(minute.title)
    run.bold = True
    run.font.size = Pt(12)

    notice = document.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = notice.add_run("BORRADOR GENERADO AUTOMÁTICAMENTE — REQUIERE REVISIÓN HUMANA")
    run.bold = True
    run.font.size = Pt(8)

    metadata_table = document.add_table(rows=0, cols=2)
    metadata_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata_table.style = "Table Grid"

    metadata_rows = [
        ("Empresa", company_name),
        ("Proyecto", metadata.get("project") or "No indicado"),
        ("Tipo de reunión", metadata.get("meeting_type") or "No indicado"),
        ("Fecha", metadata.get("meeting_date") or "No indicada"),
        ("Archivo fuente", metadata.get("source_file") or "No indicado"),
        ("Modelo IA", metadata.get("model") or "No indicado"),
    ]
    for label, value in metadata_rows:
        cells = metadata_table.add_row().cells
        _set_cell_text(cells[0], label, bold=True)
        _set_cell_shading(cells[0], "D9E2F3")
        _set_cell_text(cells[1], str(value))

    _add_heading(document, "1. Objetivo")
    if minute.objective:
        document.add_paragraph(minute.objective)
    else:
        _add_empty_message(document, "No fue posible determinar el objetivo con certeza.")

    _add_heading(document, "2. Resumen ejecutivo")
    document.add_paragraph(minute.executive_summary)

    _add_heading(document, "3. Temas tratados")
    if minute.topics:
        for topic in minute.topics:
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(f"{topic.title}: ")
            run.bold = True
            paragraph.add_run(topic.summary)
            if topic.evidence:
                paragraph.add_run(f" [Evidencia: {', '.join(topic.evidence)}]")
    else:
        _add_empty_message(document, "No se identificaron temas relevantes.")

    _add_heading(document, "4. Decisiones y acuerdos")
    if minute.decisions:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["N.º", "Decisión / acuerdo", "Evidencia"]
        for index, header in enumerate(headers):
            _set_cell_text(table.rows[0].cells[index], header, bold=True)
            _set_cell_shading(table.rows[0].cells[index], "D9E2F3")
        _set_repeat_table_header(table.rows[0])
        for index, item in enumerate(minute.decisions, start=1):
            cells = table.add_row().cells
            _set_cell_text(cells[0], str(index))
            _set_cell_text(cells[1], item.description)
            _set_cell_text(cells[2], item.evidence or "Por confirmar")
    else:
        _add_empty_message(document, "No se identificaron decisiones explícitas.")

    _add_heading(document, "5. Compromisos")
    if minute.action_items:
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = [
            "N.º",
            "Compromiso",
            "Responsable",
            "Plazo",
            "Evidencia",
            "Confianza",
        ]
        for index, header in enumerate(headers):
            _set_cell_text(table.rows[0].cells[index], header, bold=True, size=8)
            _set_cell_shading(table.rows[0].cells[index], "D9E2F3")
        _set_repeat_table_header(table.rows[0])
        for index, item in enumerate(minute.action_items, start=1):
            cells = table.add_row().cells
            values = [
                str(index),
                item.description,
                item.owner or "Por confirmar",
                item.due_date_text or item.due_date_iso or "Por confirmar",
                item.evidence or "Por confirmar",
                f"{item.confidence * 100:.0f} %",
            ]
            for cell, value in zip(cells, values, strict=False):
                _set_cell_text(cell, value, size=8)
    else:
        _add_empty_message(document, "No se identificaron compromisos explícitos.")

    _add_heading(document, "6. Pendientes")
    if minute.pending_items:
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["N.º", "Pendiente", "Responsable", "Evidencia"]
        for index, header in enumerate(headers):
            _set_cell_text(table.rows[0].cells[index], header, bold=True)
            _set_cell_shading(table.rows[0].cells[index], "D9E2F3")
        _set_repeat_table_header(table.rows[0])
        for index, item in enumerate(minute.pending_items, start=1):
            cells = table.add_row().cells
            _set_cell_text(cells[0], str(index))
            _set_cell_text(cells[1], item.description)
            _set_cell_text(cells[2], item.owner or "Por confirmar")
            _set_cell_text(cells[3], item.evidence or "Por confirmar")
    else:
        _add_empty_message(document, "No se identificaron pendientes.")

    _add_heading(document, "7. Próxima reunión")
    if minute.next_meeting:
        fields = [
            minute.next_meeting.description,
            minute.next_meeting.date_text,
            minute.next_meeting.time_text,
        ]
        description = " — ".join(item for item in fields if item)
        document.add_paragraph(description or "Información incompleta.")
        if minute.next_meeting.evidence:
            document.add_paragraph(f"Evidencia: {minute.next_meeting.evidence}")
    else:
        _add_empty_message(document, "No se indicó una próxima reunión.")

    _add_heading(document, "8. Advertencias para revisión")
    if minute.warnings:
        for warning in minute.warnings:
            document.add_paragraph(warning, style="List Bullet")
    else:
        document.add_paragraph(
            "No se generaron advertencias automáticas. "
            "De todas formas, revise la minuta antes de distribuirla."
        )

    document.add_paragraph()
    closing = document.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run(
        "Documento generado desde una transcripción. "
        "La aprobación final es responsabilidad del revisor."
    )
    run.italic = True
    run.font.size = Pt(8)

    document.save(str(path))
    return path
