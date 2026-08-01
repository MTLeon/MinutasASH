from __future__ import annotations

from copy import deepcopy
from datetime import datetime
import hashlib
from pathlib import Path
import re
import shutil
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import _Row, Table
from docx.text.paragraph import Paragraph

from src.catalog_models import TemplateManifest, TemplateValidation
from src.document_validator import validate_generated_docx
from src.models import MeetingMetadata, MinuteAnalysis
from src.runtime_paths import templates_dir


SCALAR_MARKERS: dict[str, str] = {
    "NUMERO_MINUTA": "minute_number",
    "FECHA_DOCUMENTO": "document_date",
    "FECHA_REUNION": "meeting_date",
    "LUGAR": "location",
    "MATERIA": "matter",
    "CODIGO_PROYECTO": "project_code",
    "DESCRIPCION_PROYECTO": "project_description",
    "CLIENTE": "client",
    "TOMADA_POR": "minute_taker",
    "FECHA_ELABORACION": "minute_taker_date",
    "APROBADA_POR": "approved_by",
    "FECHA_APROBACION": "approval_date",
    "TIPO_REUNION": "meeting_type",
    "VERSION_PLANTILLA": "template_version",
}
TABLE_MARKERS = {"TABLA_ASISTENTES", "TABLA_ACUERDOS"}
KNOWN_MARKERS = set(SCALAR_MARKERS) | TABLE_MARKERS
DEFAULT_REQUIRED = {
    "NUMERO_MINUTA",
    "FECHA_REUNION",
    "MATERIA",
    "CODIGO_PROYECTO",
    "CLIENTE",
    "TABLA_ASISTENTES",
    "TABLA_ACUERDOS",
}
MARKER_RE = re.compile(r"\{\{\s*([A-Z0-9_]+)\s*\}\}")


def sha256_file(path: str | Path) -> str:
    digest = hashlib.sha256()
    with Path(path).open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _iter_paragraphs(container) -> Iterable[Paragraph]:
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def _all_story_parts(document: DocumentObject):
    yield document
    for section in document.sections:
        yield section.header
        yield section.footer
        yield section.first_page_header
        yield section.first_page_footer
        yield section.even_page_header
        yield section.even_page_footer


def _all_text(document: DocumentObject) -> str:
    blocks: list[str] = []
    for part in _all_story_parts(document):
        for paragraph in _iter_paragraphs(part):
            if paragraph.text:
                blocks.append(paragraph.text)
    return "\n".join(blocks)


def _find_table_marker_rows(document: DocumentObject) -> dict[str, list[tuple[Table, _Row]]]:
    result: dict[str, list[tuple[Table, _Row]]] = {name: [] for name in TABLE_MARKERS}
    for part in _all_story_parts(document):
        for table in part.tables:
            for row in table.rows:
                text = " | ".join(cell.text for cell in row.cells)
                for marker in TABLE_MARKERS:
                    if f"{{{{{marker}}}}}" in text.replace(" ", ""):
                        result[marker].append((table, row))
    return result


def validate_template(
    path: str | Path,
    required_markers: Iterable[str] | None = None,
) -> TemplateValidation:
    template_path = Path(path)
    if not template_path.is_file() or template_path.suffix.lower() != ".docx":
        return TemplateValidation(
            valid=False,
            warnings=["El archivo debe existir y utilizar extensión .docx."],
        )
    try:
        document = Document(template_path)
    except Exception as exc:
        return TemplateValidation(valid=False, warnings=[f"Word no pudo abrir la plantilla: {exc}"])

    text = _all_text(document)
    found = sorted(set(MARKER_RE.findall(text)))
    required = set(required_markers or DEFAULT_REQUIRED)
    missing = sorted(required - set(found))
    unknown = sorted(set(found) - KNOWN_MARKERS)
    marker_rows = _find_table_marker_rows(document)
    table_counts = {key: len(value) for key, value in marker_rows.items()}
    warnings: list[str] = []
    for marker in TABLE_MARKERS:
        if marker in found and table_counts.get(marker, 0) != 1:
            warnings.append(
                f"El marcador {marker} debe aparecer exactamente una vez dentro de una fila de tabla."
            )
    if len(document.sections) == 0:
        warnings.append("La plantilla no contiene una sección de página válida.")
    if not document.tables:
        warnings.append("La plantilla no contiene tablas; revise el formato corporativo.")
    valid = not missing and not unknown and all(
        table_counts.get(marker, 0) == 1 for marker in TABLE_MARKERS
    )
    return TemplateValidation(
        valid=valid,
        markers_found=found,
        missing_required=missing,
        unknown_markers=unknown,
        warnings=warnings,
        table_markers=table_counts,
    )


def install_template_file(
    source_path: str | Path,
    manifest: TemplateManifest,
) -> Path:
    source = Path(source_path)
    validation = validate_template(source, manifest.required_markers or None)
    if not validation.valid:
        problems = []
        if validation.missing_required:
            problems.append("Faltan: " + ", ".join(validation.missing_required))
        if validation.unknown_markers:
            problems.append("Desconocidos: " + ", ".join(validation.unknown_markers))
        problems.extend(validation.warnings)
        raise ValueError("La plantilla no cumple el contrato documental. " + " | ".join(problems))
    destination = templates_dir() / manifest.template_key / f"{manifest.version_label}.docx"
    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_suffix(".docx.tmp")
    shutil.copy2(source, temporary)
    temporary.replace(destination)
    return destination


def _replace_paragraph_text(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    original = paragraph.text
    updated = original
    for marker, value in replacements.items():
        updated = updated.replace(f"{{{{{marker}}}}}", value)
        updated = re.sub(r"\{\{\s*" + re.escape(marker) + r"\s*\}\}", value, updated)
    if updated == original:
        return
    # Mantener el formato base del primer run. Los marcadores pueden estar
    # fragmentados en varios runs y por eso se reconstruye el párrafo.
    first = paragraph.runs[0] if paragraph.runs else None
    properties = deepcopy(first._r.rPr) if first is not None and first._r.rPr is not None else None
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    run = paragraph.add_run(updated)
    if properties is not None:
        run._r.insert(0, properties)


def _fill_cell(cell, value: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.add_run(value or "")


def _clone_row_before(table: Table, template_row: _Row, values: list[str]) -> None:
    new_tr = deepcopy(template_row._tr)
    template_row._tr.addprevious(new_tr)
    row = _Row(new_tr, table)
    cells = row.cells
    for index, cell in enumerate(cells):
        _fill_cell(cell, values[index] if index < len(values) else "")


def _attendee_values(metadata: MeetingMetadata, column_count: int) -> list[list[str]]:
    rows: list[list[str]] = []
    for index, attendee in enumerate(metadata.attendees, start=1):
        full = [
            str(index),
            attendee.initials or "",
            attendee.name,
            attendee.email or "",
            attendee.role or "",
            attendee.organization or "",
        ]
        if column_count >= 6:
            rows.append(full[:column_count])
        elif column_count == 5:
            rows.append([full[0], full[2], full[3], full[4], full[5]])
        elif column_count == 4:
            rows.append([full[0], full[2], full[4], full[5]])
        elif column_count == 3:
            rows.append([full[0], full[2], full[5]])
        elif column_count == 2:
            rows.append([full[2], full[5]])
        else:
            rows.append([f"{full[0]}. {full[2]} - {full[5]}"])
    if not rows:
        rows.append(["Sin participantes confirmados"] + [""] * max(0, column_count - 1))
    return rows


def _item_values(analysis: MinuteAnalysis, column_count: int) -> list[list[str]]:
    rows: list[list[str]] = []
    active = [item for item in analysis.items if item.review_status != "descartado"]
    for index, item in enumerate(active, start=1):
        responsible = item.responsible or ("Informativo" if item.category in {"informativo", "acuerdo"} else "Por confirmar")
        due = item.due_date_text or item.due_date_iso or ("N.A." if item.category in {"informativo", "acuerdo", "pendiente"} else "Por confirmar")
        description = item.description
        if item.project_code and not description.casefold().startswith(f"proyecto {item.project_code}".casefold()):
            description = f"Proyecto {item.project_code} — {description}"
        if item.category == "pendiente" and not description.lower().startswith("pendiente"):
            description = "Pendiente: " + description
        full = [str(index), description, responsible, due]
        if column_count >= 4:
            rows.append(full + [""] * (column_count - 4))
        elif column_count == 3:
            rows.append([full[0], full[1], full[2]])
        elif column_count == 2:
            rows.append([full[1], f"{full[2]} / {full[3]}"])
        else:
            rows.append([f"{full[0]}. {full[1]} | {full[2]} | {full[3]}"])
    if not rows:
        rows.append(["Sin acuerdos o compromisos identificados"] + [""] * max(0, column_count - 1))
    return rows


def _replace_table_marker(document: DocumentObject, marker: str, rows: list[list[str]]) -> None:
    matches = _find_table_marker_rows(document).get(marker, [])
    if len(matches) != 1:
        raise ValueError(f"La plantilla debe contener una fila única con {{{{{marker}}}}}.")
    table, placeholder = matches[0]
    column_count = len(placeholder.cells)
    for values in rows:
        normalized = values[:column_count] + [""] * max(0, column_count - len(values))
        _clone_row_before(table, placeholder, normalized)
    table._tbl.remove(placeholder._tr)


def _format_date(value: str | None) -> str:
    if not value:
        return ""
    try:
        return datetime.strptime(value, "%Y-%m-%d").strftime("%d-%m-%Y")
    except ValueError:
        return value


def render_template_document(
    template_path: str | Path,
    metadata: MeetingMetadata,
    analysis: MinuteAnalysis,
    output_path: str | Path,
) -> Path:
    validation = validate_template(template_path)
    if not validation.valid:
        raise ValueError(
            "La plantilla seleccionada dejó de ser válida: "
            + "; ".join(validation.missing_required + validation.unknown_markers + validation.warnings)
        )
    document = Document(template_path)
    values = metadata.model_dump()
    values["document_date"] = _format_date(metadata.document_date)
    values["meeting_date"] = _format_date(metadata.meeting_date)
    values["minute_taker_date"] = _format_date(metadata.minute_taker_date)
    values["approval_date"] = _format_date(metadata.approval_date)
    replacements = {
        marker: str(values.get(field) or "")
        for marker, field in SCALAR_MARKERS.items()
    }
    for part in _all_story_parts(document):
        for paragraph in _iter_paragraphs(part):
            _replace_paragraph_text(paragraph, replacements)

    attendee_rows = _attendee_values(metadata, len(_find_table_marker_rows(document)["TABLA_ASISTENTES"][0][1].cells))
    item_rows = _item_values(analysis, len(_find_table_marker_rows(document)["TABLA_ACUERDOS"][0][1].cells))
    _replace_table_marker(document, "TABLA_ASISTENTES", attendee_rows)
    _replace_table_marker(document, "TABLA_ACUERDOS", item_rows)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = metadata.minute_number or "Minuta de reunión"
    document.core_properties.subject = metadata.matter or "Minuta de reunión"
    document.core_properties.author = metadata.minute_taker or "ASH Ingeniería y Proyectos"
    document.core_properties.comments = (
        f"Plantilla {metadata.template_key or 'administrada'} "
        f"versión {metadata.template_version or 'sin versión'}"
    )
    document.save(output)
    validate_generated_docx(output, metadata, analysis, minimum_tables=2)
    return output


def create_test_metadata() -> tuple[MeetingMetadata, MinuteAnalysis]:
    from src.models import Attendee, MeetingItem

    metadata = MeetingMetadata(
        meeting_type="cliente",
        minute_number="P0000-MRE-PR-00",
        document_date="2026-07-31",
        meeting_date="2026-07-31",
        location="Microsoft Teams",
        matter="Validación de plantilla documental",
        project_code="P0000",
        project_description="Proyecto de prueba",
        client="Cliente de prueba",
        minute_taker="Usuario de prueba",
        minute_taker_date="2026-07-31",
        approved_by="Aprobador de prueba",
        approval_date="2026-07-31",
        attendees=[
            Attendee(initials="UP", name="Usuario de prueba", email="usuario@ash.cl", role="Ingeniero", organization="ASH"),
            Attendee(initials="CP", name="Contacto de prueba", email="contacto@cliente.cl", role="Especialista", organization="Cliente"),
        ],
    )
    analysis = MinuteAnalysis(
        objective="Validar el formato documental.",
        executive_summary="Se verificó la construcción de la plantilla.",
        items=[
            MeetingItem(category="informativo", description="Se presentó el alcance de la reunión.", review_status="aprobado", origin="manual", confidence=1.0),
            MeetingItem(category="compromiso", description="ASH enviará los planos actualizados.", responsible="Usuario de prueba", due_date_text="viernes 7 de agosto", review_status="aprobado", origin="manual", confidence=1.0),
            MeetingItem(category="pendiente", description="Confirmar el número de señales analógicas.", responsible="Contacto de prueba", review_status="aprobado", origin="manual", confidence=1.0),
        ],
    )
    return metadata, analysis
