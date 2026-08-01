from __future__ import annotations

"""Fuentes flexibles para reuniones.

La aplicación mantiene el VTT de Teams como fuente preferida, pero puede
normalizar transcripciones Word, texto copiado y notas manuales sin alterar el
pipeline de análisis. Todos los formatos se convierten en ``TranscriptSegment``
antes de llegar al modelo.
"""

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import re
from typing import Literal

from docx import Document

from src.runtime_paths import drafts_dir
from src.storage import safe_component
from src.vtt_reader import (
    TranscriptSegment,
    is_valid_speaker_name,
    merge_adjacent_segments,
    read_teams_vtt,
)


SourceType = Literal["vtt", "docx", "txt", "pasted", "notes"]
SourceQuality = Literal["alta", "media", "baja"]

SOURCE_TYPE_LABELS: dict[str, str] = {
    "vtt": "Transcripción de Teams (VTT)",
    "docx": "Transcripción o notas Word",
    "txt": "Archivo de texto",
    "pasted": "Conversación pegada",
    "notes": "Notas manuales",
}

SOURCE_QUALITY_LABELS: dict[str, str] = {
    "alta": "Alta · conserva hablantes y tiempos",
    "media": "Media · requiere revisión reforzada",
    "baja": "Baja · revisión completa obligatoria",
}

_TIMESTAMP_SPEAKER_RE = re.compile(
    r"^\s*\[?(?P<timestamp>\d{1,2}:\d{2}(?::\d{2})?(?:[.,]\d{1,3})?)\]?\s*"
    r"[-–—|]?\s*(?P<speaker>[^:\n]{2,100}):\s*(?P<text>.+?)\s*$"
)
_SPEAKER_RE = re.compile(r"^\s*(?P<speaker>[^:\n]{2,100}):\s*(?P<text>.+?)\s*$")
_TIMESTAMP_ONLY_RE = re.compile(
    r"^\s*\[?(?P<timestamp>\d{1,2}:\d{2}(?::\d{2})?(?:[.,]\d{1,3})?)\]?\s*$"
)


@dataclass(frozen=True)
class MeetingSource:
    path: Path
    source_type: SourceType
    quality: SourceQuality
    segments: list[TranscriptSegment]
    warnings: tuple[str, ...] = ()

    @property
    def display_name(self) -> str:
        return SOURCE_TYPE_LABELS.get(self.source_type, self.source_type)


def _normalize_timestamp(value: str | None, index: int) -> str:
    if not value:
        seconds = index * 5
        return f"{seconds // 3600:02d}:{(seconds % 3600) // 60:02d}:{seconds % 60:02d}.000"
    text = value.replace(",", ".").strip("[] ")
    parts = text.split(":")
    if len(parts) == 2:
        text = f"00:{text}"
    if "." not in text:
        text += ".000"
    else:
        head, fraction = text.split(".", 1)
        text = f"{head}.{fraction[:3].ljust(3, '0')}"
    return text


def _plus_seconds(timestamp: str, seconds: float = 4.0) -> str:
    hours, minutes, second_text = timestamp.split(":")
    total = int(hours) * 3600 + int(minutes) * 60 + float(second_text) + seconds
    return f"{int(total // 3600):02d}:{int((total % 3600) // 60):02d}:{total % 60:06.3f}"


def _looks_like_speaker(value: str) -> bool:
    text = " ".join(value.split()).strip()
    if not text or len(text) > 80 or not is_valid_speaker_name(text):
        return False
    lowered = text.casefold()
    if lowered.startswith(("tema ", "acuerdo", "compromiso", "pendiente", "próxima reunión")):
        return False
    words = text.split()
    # Nombres completos, roles breves o etiquetas como "Cliente".
    return len(words) <= 8 and not text.endswith((".", "?", "!"))


def parse_text_transcript(
    text: str,
    default_speaker: str = "Notas de reunión",
) -> list[TranscriptSegment]:
    """Convierte texto libre en segmentos normalizados.

    Admite líneas ``[00:01:20] Nombre: texto``, ``Nombre: texto`` y párrafos
    sin hablante. Las continuaciones se anexan al último segmento para no perder
    contenido cuando Teams copia una intervención en varias líneas.
    """

    lines = [line.rstrip() for line in str(text or "").replace("\r\n", "\n").split("\n")]
    segments: list[TranscriptSegment] = []
    pending_timestamp: str | None = None

    def append_segment(speaker: str, body: str, timestamp: str | None = None) -> None:
        cleaned = re.sub(r"\s+", " ", body).strip()
        if not cleaned:
            return
        start = _normalize_timestamp(timestamp, len(segments))
        segments.append(
            TranscriptSegment(
                start=start,
                end=_plus_seconds(start),
                speaker=" ".join(speaker.split()).strip() or default_speaker,
                text=cleaned,
            )
        )

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        timestamp_only = _TIMESTAMP_ONLY_RE.match(line)
        if timestamp_only:
            pending_timestamp = timestamp_only.group("timestamp")
            continue
        match = _TIMESTAMP_SPEAKER_RE.match(line)
        if match and _looks_like_speaker(match.group("speaker")):
            append_segment(match.group("speaker"), match.group("text"), match.group("timestamp"))
            pending_timestamp = None
            continue
        match = _SPEAKER_RE.match(line)
        if match and _looks_like_speaker(match.group("speaker")):
            append_segment(match.group("speaker"), match.group("text"), pending_timestamp)
            pending_timestamp = None
            continue
        if segments:
            previous = segments[-1]
            segments[-1] = TranscriptSegment(
                previous.start,
                previous.end,
                previous.speaker,
                f"{previous.text} {line}".strip(),
            )
        else:
            append_segment(default_speaker, line, pending_timestamp)
            pending_timestamp = None

    if not segments:
        raise ValueError("La fuente no contiene texto utilizable para generar una minuta.")
    return merge_adjacent_segments(segments, maximum_gap_seconds=6.0)


def _read_docx_text(path: Path) -> str:
    document = Document(path)
    lines: list[str] = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            lines.append(value)
    for table in document.tables:
        for row in table.rows:
            values = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
            values = [value for value in values if value]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)


def infer_source_type(path: str | Path, preferred: str | None = None) -> SourceType:
    if preferred in SOURCE_TYPE_LABELS:
        return preferred  # type: ignore[return-value]
    suffix = Path(path).suffix.casefold()
    mapping = {".vtt": "vtt", ".docx": "docx", ".txt": "txt"}
    if suffix not in mapping:
        raise ValueError("Formato no admitido. Use VTT, TXT o DOCX.")
    return mapping[suffix]  # type: ignore[return-value]


def source_quality(source_type: str, segments: list[TranscriptSegment]) -> SourceQuality:
    if source_type == "vtt":
        return "alta"
    identified = sum(1 for segment in segments if segment.speaker not in {"Notas de reunión", "Hablante no identificado"})
    ratio = identified / max(len(segments), 1)
    if source_type == "docx" and ratio >= 0.6:
        return "alta"
    if ratio >= 0.25 or source_type in {"docx", "txt", "pasted", "notes"}:
        return "media"
    return "baja"


def read_meeting_source(path: str | Path, preferred_type: str | None = None) -> MeetingSource:
    source_path = Path(path).expanduser().resolve()
    if not source_path.is_file():
        raise FileNotFoundError(f"No existe la fuente de reunión: {source_path}")
    source_type = infer_source_type(source_path, preferred_type)
    warnings: list[str] = []
    if source_type == "vtt":
        segments = read_teams_vtt(source_path, merge_adjacent=False)
    elif source_type == "docx":
        segments = parse_text_transcript(_read_docx_text(source_path))
        warnings.append("El documento Word no siempre conserva marcas de tiempo o atribución completa de hablantes.")
    else:
        segments = parse_text_transcript(source_path.read_text(encoding="utf-8-sig", errors="replace"))
        warnings.append("La fuente textual requiere confirmar hablantes, fechas y contexto durante la revisión.")
    quality = source_quality(source_type, segments)
    return MeetingSource(source_path, source_type, quality, segments, tuple(warnings))


def create_text_source(
    text: str,
    source_type: SourceType = "pasted",
    suggested_name: str = "reunion",
) -> MeetingSource:
    if source_type not in {"pasted", "notes", "txt"}:
        raise ValueError("La captura manual debe ser texto pegado o notas.")
    content = str(text or "").strip()
    if len(content) < 8:
        raise ValueError("Ingrese contenido suficiente para preparar la minuta.")
    drafts_dir().mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    destination = drafts_dir() / f"{safe_component(suggested_name, 'reunion')}_{source_type}_{stamp}.txt"
    destination.write_text(content, encoding="utf-8")
    segments = parse_text_transcript(content)
    return MeetingSource(
        destination,
        source_type,
        source_quality(source_type, segments),
        segments,
        ("Fuente creada manualmente dentro de Minutas ASH; revisión completa recomendada.",),
    )


def supported_filetypes() -> list[tuple[str, str]]:
    return [
        ("Fuentes de reunión", "*.vtt *.txt *.docx"),
        ("Transcripción de Teams", "*.vtt"),
        ("Texto", "*.txt"),
        ("Documento Word", "*.docx"),
        ("Todos los archivos", "*.*"),
    ]
