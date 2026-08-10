from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.metadata import format_date
from src.models import Attendee, MeetingItem, MeetingMetadata, MinuteAnalysis

BLUE = "1F497D"
FONT = "Segoe UI"


def _set_cell_margins(cell, top=45, start=70, bottom=45, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str = BLUE, outer_size: int = 6, inner_size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(outer_size if edge in {"top", "left", "bottom", "right"} else inner_size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(inches)


def _set_table_grid_widths(table, widths: list[float]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            _set_cell_width(row.cells[idx], width)


def _set_row_repeat(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_update_fields(document: DocxDocument) -> None:
    settings = document.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def _add_field(paragraph, instruction: str, placeholder: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def _style_run(run, size: float = 9, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _prepare_paragraph(paragraph, align=None, before=0, after=0, line=1.0) -> None:
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def _clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _prepare_paragraph(paragraph)


def _write_cell(cell, text: str, bold=False, size=9, align=None) -> None:
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text or "")
    _style_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)


def _format_attendee(attendee: Attendee) -> str:
    prefix = f"({attendee.initials}) " if attendee.initials else ""
    parts = [f"{prefix}{attendee.name}"]
    if attendee.email:
        parts.append(f" / {attendee.email}")
    if attendee.role:
        parts.append(f" - {attendee.role}")
    return "".join(parts)


def _responsible_and_date(item: MeetingItem) -> tuple[str, str]:
    if item.category in {"informativo", "acuerdo"} and not item.responsible:
        date_value = item.due_date_text or item.due_date_iso or "N.A."
        return "Informativo", date_value
    responsible = item.responsible or "Por confirmar"
    date_value = item.due_date_text or item.due_date_iso or (
        "Por confirmar" if item.category == "compromiso" else "N.A."
    )
    return responsible, date_value


def _add_description(cell, item: MeetingItem) -> None:
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    _prepare_paragraph(paragraph, line=1.0)
    if item.title:
        title_run = paragraph.add_run(f"{item.title}:\n")
        _style_run(title_run, size=9, bold=True)
    if item.source_speaker:
        speaker_run = paragraph.add_run(f"{item.source_speaker} ")
        _style_run(speaker_run, size=9, bold=True)
    description = item.description
    if item.project_code and not description.casefold().startswith(f"proyecto {item.project_code}".casefold()):
        description = f"Proyecto {item.project_code} — {description}"
    body = paragraph.add_run(description)
    _style_run(body, size=9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell, top=60, bottom=60)


def _configure_section(document: DocxDocument) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(1.38)
    section.bottom_margin = Inches(0.55)
    section.header_distance = Inches(0.12)
    section.footer_distance = Inches(0.25)


def _build_header(document: DocxDocument, metadata: MeetingMetadata, logo_path: Path, border_color: str) -> None:
    header = document.sections[0].header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    _prepare_paragraph(header_paragraph, after=0)

    table = header.add_table(rows=4, cols=4, width=Inches(7.4))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, border_color, outer_size=6, inner_size=4)

    widths = [1.25, 4.3, 0.55, 1.30]
    _set_table_grid_widths(table, widths)

    logo_cell = table.cell(0, 0).merge(table.cell(3, 0))
    title_cell = table.cell(0, 1).merge(table.cell(3, 1))
    page_cell = table.cell(3, 2).merge(table.cell(3, 3))

    _clear_cell(logo_cell)
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path.exists():
        logo_p.add_run().add_picture(str(logo_path), width=Inches(1.12))
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(logo_cell, top=20, bottom=20)

    _clear_cell(title_cell)
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = title_p.add_run("SISTEMA DE GESTION INTEGRADO\n\n")
    _style_run(r1, size=8)
    r2 = title_p.add_run("Minuta de Reunión")
    _style_run(r2, size=16, bold=True)
    title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    _write_cell(table.cell(0, 2), "N°:", bold=True, size=9)
    _write_cell(table.cell(0, 3), metadata.minute_number or "Por asignar", size=9)
    _write_cell(table.cell(1, 2), "Fecha:", bold=True, size=9)
    _write_cell(table.cell(1, 3), format_date(metadata.document_date, "-"), size=9)
    _write_cell(table.cell(2, 2), "", size=8)
    _write_cell(table.cell(2, 3), "", size=8)

    _clear_cell(page_cell)
    page_p = page_cell.paragraphs[0]
    page_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = page_p.add_run("Pág. ")
    _style_run(r, size=9, bold=True)
    _add_field(page_p, "PAGE", "1")
    r = page_p.add_run(" de ")
    _style_run(r, size=9, bold=True)
    _add_field(page_p, "NUMPAGES", "1")
    for run in page_p.runs:
        _style_run(run, size=9, bold=True)
    page_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # Remove the empty paragraph that precedes the table as much as Word allows.
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_paragraph.paragraph_format.line_spacing = 0.2


def _build_general_data(document: DocxDocument, metadata: MeetingMetadata, border_color: str) -> None:
    table = document.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, border_color, outer_size=6, inner_size=4)
    widths = [0.88, 1.20, 0.70, 4.62]
    _set_table_grid_widths(table, widths)

    _write_cell(table.cell(0, 0), "Fecha:", bold=True)
    _write_cell(table.cell(0, 1), format_date(metadata.meeting_date, "/"))
    _write_cell(table.cell(0, 2), "Lugar:", bold=True)
    _write_cell(table.cell(0, 3), metadata.location or "Por confirmar")

    matter_cell = table.cell(1, 1).merge(table.cell(1, 3))
    _write_cell(table.cell(1, 0), "Materia:", bold=True)
    _write_cell(matter_cell, metadata.matter or "Por confirmar")

    project_desc = table.cell(2, 2).merge(table.cell(2, 3))
    _write_cell(table.cell(2, 0), "Proyecto:", bold=True)
    _write_cell(table.cell(2, 1), metadata.project_code or "Por confirmar")
    _write_cell(project_desc, metadata.project_description or "Por confirmar")

    client_cell = table.cell(3, 1).merge(table.cell(3, 3))
    _write_cell(table.cell(3, 0), "Cliente:", bold=True)
    _write_cell(client_cell, metadata.client or "Por confirmar")

    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _build_attendees(document: DocxDocument, attendees: list[Attendee], border_color: str) -> None:
    count = max(len(attendees), 1)
    table = document.add_table(rows=2 + count, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, border_color, outer_size=4, inner_size=4)
    widths = [0.58, 5.48, 1.34]
    _set_table_grid_widths(table, widths)

    title = table.cell(0, 0).merge(table.cell(0, 2))
    _write_cell(title, "Asistentes", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    headers = ["Id", "Nombre", "Organización"]
    for idx, header in enumerate(headers):
        _write_cell(table.cell(1, idx), header, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_row_repeat(table.rows[0])
    _set_row_repeat(table.rows[1])

    if attendees:
        for row_idx, attendee in enumerate(attendees, start=2):
            _write_cell(table.cell(row_idx, 0), str(attendee.id or row_idx - 1), align=WD_ALIGN_PARAGRAPH.CENTER)
            _write_cell(table.cell(row_idx, 1), _format_attendee(attendee))
            _write_cell(table.cell(row_idx, 2), attendee.organization or "Por confirmar", align=WD_ALIGN_PARAGRAPH.CENTER)
            _prevent_row_split(table.rows[row_idx])
    else:
        _write_cell(table.cell(2, 0), "1", align=WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell(table.cell(2, 1), "Asistentes por confirmar")
        _write_cell(table.cell(2, 2), "Por confirmar", align=WD_ALIGN_PARAGRAPH.CENTER)

    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _build_approval(document: DocxDocument, metadata: MeetingMetadata, border_color: str) -> None:
    table = document.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, border_color, outer_size=6, inner_size=4)
    widths = [1.65, 3.80, 0.75, 1.20]
    _set_table_grid_widths(table, widths)

    values = [
        ("Minuta Tomada por:", metadata.minute_taker or "Por confirmar", "Fecha:", format_date(metadata.minute_taker_date, "/")),
        ("Minuta Aprobada por:", metadata.approved_by or "", "Fecha:", format_date(metadata.approval_date, "/")),
    ]
    for row_idx, row_values in enumerate(values):
        _write_cell(table.cell(row_idx, 0), row_values[0], bold=True)
        _write_cell(table.cell(row_idx, 1), row_values[1])
        _write_cell(table.cell(row_idx, 2), row_values[2], bold=True)
        _write_cell(table.cell(row_idx, 3), row_values[3])

    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _build_items(document: DocxDocument, items: list[MeetingItem], border_color: str) -> None:
    items = [
        item for item in items
        if getattr(item, "review_status", "pendiente") != "descartado"
    ]
    count = max(len(items), 1)
    table = document.add_table(rows=2 + count, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders(table, border_color, outer_size=6, inner_size=4)
    widths = [0.58, 4.70, 1.20, 0.92]
    _set_table_grid_widths(table, widths)

    title = table.cell(0, 0).merge(table.cell(0, 3))
    _write_cell(title, "Acuerdos y Compromisos", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    headers = ["N°", "Descripción", "Responsable", "Fecha"]
    for idx, header in enumerate(headers):
        _write_cell(table.cell(1, idx), header, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_row_repeat(table.rows[0])
    _set_row_repeat(table.rows[1])

    if items:
        for row_idx, item in enumerate(items, start=2):
            responsible, due_date = _responsible_and_date(item)
            _write_cell(table.cell(row_idx, 0), str(row_idx - 1), align=WD_ALIGN_PARAGRAPH.CENTER)
            _add_description(table.cell(row_idx, 1), item)
            _write_cell(table.cell(row_idx, 2), responsible, align=WD_ALIGN_PARAGRAPH.CENTER)
            _write_cell(table.cell(row_idx, 3), due_date, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        _write_cell(table.cell(2, 0), "1", align=WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell(table.cell(2, 1), "Sin acuerdos o compromisos identificados.")
        _write_cell(table.cell(2, 2), "Informativo", align=WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell(table.cell(2, 3), "N.A.", align=WD_ALIGN_PARAGRAPH.CENTER)


def generate_ash_docx(
    analysis: MinuteAnalysis,
    metadata: MeetingMetadata,
    output_path: str | Path,
    logo_path: str | Path,
    border_color: str = BLUE,
) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_section(document)
    _set_update_fields(document)

    styles = document.styles
    styles["Normal"].font.name = FONT
    styles["Normal"].font.size = Pt(9)

    _build_header(document, metadata, Path(logo_path), border_color)
    _build_general_data(document, metadata, border_color)
    _build_attendees(document, metadata.attendees, border_color)
    _build_approval(document, metadata, border_color)
    _build_items(document, analysis.items, border_color)

    document.core_properties.title = metadata.minute_number or "Minuta de Reunión"
    document.core_properties.subject = metadata.matter or "Minuta ASH"
    document.core_properties.author = "ASH Ingeniería y Proyectos"
    document.core_properties.last_modified_by = metadata.minute_taker or "ASH Ingeniería y Proyectos"
    document.core_properties.comments = "Documento emitido mediante Minutas ASH."
    document.save(str(output))
    return output
