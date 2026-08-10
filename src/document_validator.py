from __future__ import annotations

from pathlib import Path

from docx import Document

from src.models import MeetingMetadata, MinuteAnalysis


class DocumentValidationError(RuntimeError):
    pass


def validate_generated_docx(
    path: str | Path,
    metadata: MeetingMetadata,
    analysis: MinuteAnalysis,
    *,
    minimum_tables: int = 4,
) -> None:
    """Verificación estructural mínima antes de publicar el documento."""
    document_path = Path(path)
    if not document_path.is_file() or document_path.stat().st_size < 10_000:
        raise DocumentValidationError("El documento generado está vacío o incompleto.")

    try:
        document = Document(str(document_path))
    except Exception as exc:
        raise DocumentValidationError("El documento generado no pudo volver a abrirse.") from exc

    total_table_count = len(document.tables) + sum(
        len(section.header.tables) + len(section.footer.tables)
        for section in document.sections
    )
    if total_table_count < minimum_tables:
        raise DocumentValidationError(
            "El documento no contiene las tablas corporativas mínimas requeridas."
        )

    tables = list(document.tables)
    for section in document.sections:
        tables.extend(section.header.tables)
        tables.extend(section.footer.tables)
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]
    for section in document.sections:
        paragraph_text.extend(paragraph.text for paragraph in section.header.paragraphs)
        paragraph_text.extend(paragraph.text for paragraph in section.footer.paragraphs)
    all_text = "\n".join(paragraph_text + [
        cell.text
        for table in tables
        for row in table.rows
        for cell in row.cells
    ])
    required = [
        "Acuerdos y Compromisos",
        "Asistentes",
        metadata.minute_number or "",
    ]
    for value in required:
        if value and value not in all_text:
            raise DocumentValidationError(
                f"El documento no contiene el campo requerido: {value}"
            )

    if analysis.items:
        first_description = analysis.items[0].description[:60]
        if first_description and first_description not in all_text:
            raise DocumentValidationError(
                "El documento no contiene los puntos revisados de la minuta."
            )
