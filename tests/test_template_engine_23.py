from __future__ import annotations

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document

from src.catalog_models import TemplateManifest
from src.database import AppDatabase
from src.documents.registry import get_document_provider
from src.template_engine import (
    create_test_metadata,
    render_template_document,
    sha256_file,
    validate_template,
)
from src.template_service import TemplateService

ROOT = Path(__file__).resolve().parents[1]
SAMPLE = ROOT / "plantillas" / "Plantilla_Marcadores_ASH_2.3.docx"


class TemplateEngine23Tests(unittest.TestCase):
    def test_sample_template_validates_and_renders(self) -> None:
        result = validate_template(SAMPLE)
        self.assertTrue(result.valid, result.model_dump())
        metadata, analysis = create_test_metadata()
        with TemporaryDirectory() as temp:
            output = Path(temp) / "test.docx"
            render_template_document(SAMPLE, metadata, analysis, output)
            self.assertTrue(output.is_file())
            document = Document(output)
            tables = list(document.tables)
            for section in document.sections:
                tables.extend(section.header.tables)
                tables.extend(section.footer.tables)
            text = "\n".join(
                cell.text
                for table in tables
                for row in table.rows
                for cell in row.cells
            )
            self.assertIn("P0000-MRE-PR-00", text)
            self.assertIn("ASH enviará los planos actualizados", text)
            self.assertNotIn("{{TABLA_ASISTENTES}}", text)

    def test_managed_document_provider_generates_from_selected_template(self) -> None:
        metadata, analysis = create_test_metadata()
        metadata.template_key = "minuta_piloto"
        metadata.template_version = "1.0"
        with TemporaryDirectory() as temp:
            output = Path(temp) / "managed.docx"
            provider = get_document_provider("managed_template_v1")
            result = provider.generate(
                analysis,
                metadata,
                output,
                {"managed_template_path": str(SAMPLE)},
            )
            self.assertEqual(result, output)
            self.assertTrue(output.is_file())
            self.assertGreater(output.stat().st_size, 10_000)

    def test_project_can_select_an_active_template_version(self) -> None:
        with TemporaryDirectory() as temp:
            db = AppDatabase(Path(temp) / "minutas.db")
            validation = validate_template(SAMPLE)
            manifest = TemplateManifest(
                template_key="minuta_proyecto",
                display_name="Minuta de proyecto",
                version_label="3.0",
                document_type="meeting_minutes",
            )
            version_id = db.register_template_version(
                manifest, validation, str(SAMPLE), sha256_file(SAMPLE), state="testing"
            )
            db.activate_template_version(version_id)
            db.upsert_project_profile(
                {
                    "code": "P1234",
                    "description": "Proyecto con formato",
                    "client": "Cliente",
                    "template_version_id": version_id,
                }
            )
            resolved = db.resolve_template_version(project_code="P1234")
            self.assertIsNotNone(resolved)
            self.assertEqual(resolved["id"], version_id)

    def test_template_requires_test_document_before_activation(self) -> None:
        with TemporaryDirectory() as temp:
            root = Path(temp)
            db = AppDatabase(root / "minutas.db")
            service = TemplateService(db)
            with (
                patch("src.template_engine.templates_dir", return_value=root / "templates"),
                patch("src.template_service.templates_dir", return_value=root / "templates"),
                patch("src.template_service.records_dir", return_value=root / "records"),
            ):
                version_id = service.install(
                    SAMPLE,
                    template_key="control_prueba",
                    display_name="Control de prueba",
                    version_label="1.0",
                )
                self.assertEqual(db.get_template_version(version_id)["state"], "draft")
                with self.assertRaises(ValueError):
                    service.activate(version_id)
                output = service.create_test_document(version_id)
                self.assertTrue(output.is_file())
                self.assertEqual(db.get_template_version(version_id)["state"], "testing")
                service.activate(version_id)
                self.assertEqual(db.get_template_version(version_id)["state"], "active")

    def test_template_version_registration_and_activation(self) -> None:
        with TemporaryDirectory() as temp:
            db = AppDatabase(Path(temp) / "minutas.db")
            validation = validate_template(SAMPLE)
            manifest = TemplateManifest(
                template_key="minuta_piloto",
                display_name="Minuta piloto",
                version_label="1.0",
                document_type="meeting_minutes",
            )
            version_id = db.register_template_version(
                manifest, validation, str(SAMPLE), sha256_file(SAMPLE), state="testing"
            )
            db.activate_template_version(version_id)
            row = db.get_template_version(version_id)
            self.assertEqual(row["state"], "active")
            self.assertEqual(row["is_active"], 1)
            resolved = db.resolve_template_version(default_template_key="minuta_piloto")
            self.assertEqual(resolved["id"], version_id)


if __name__ == "__main__":
    unittest.main()
