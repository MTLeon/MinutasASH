from pathlib import Path
import tempfile
import unittest

from src.docx_writer_ash import generate_ash_docx
from src.document_validator import validate_generated_docx
from src.models import Attendee, MeetingItem, MeetingMetadata, MinuteAnalysis


ROOT = Path(__file__).resolve().parents[1]


class DocumentGenerationTests(unittest.TestCase):
    def test_generates_valid_ash_document(self):
        metadata = MeetingMetadata(
            minute_number="DEMO-MRE-PR-00",
            document_date="2026-07-30",
            meeting_date="2026-07-30",
            matter="Reunión de coordinación",
            project_code="DEMO-001",
            project_description="Proyecto de demostración",
            client="Cliente",
            minute_taker="Ana Pérez",
            minute_taker_date="2026-07-30",
            attendees=[
                Attendee(id=1, initials="AP", name="Ana Pérez", organization="ASH")
            ],
        )
        analysis = MinuteAnalysis(
            executive_summary="Prueba",
            items=[
                MeetingItem(
                    category="compromiso",
                    description="Entregar los planos revisados.",
                    responsible="Ana Pérez",
                    due_date_iso="2026-08-03",
                    evidence="00:01:00",
                    confidence=0.95,
                )
            ],
        )
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "minuta.docx"
            generate_ash_docx(
                analysis,
                metadata,
                output,
                ROOT / "assets" / "logo_ash.png",
            )
            validate_generated_docx(output, metadata, analysis)
            self.assertTrue(output.is_file())

    def test_discarded_items_are_not_written(self):
        from docx import Document

        metadata = MeetingMetadata(
            minute_number="DEMO-MRE-PR-01",
            document_date="2026-07-30",
            meeting_date="2026-07-30",
            matter="Reunión de coordinación",
            project_code="DEMO-001",
            project_description="Proyecto de demostración",
            client="Cliente",
            minute_taker="Ana Pérez",
            minute_taker_date="2026-07-30",
            attendees=[Attendee(id=1, initials="AP", name="Ana Pérez", organization="ASH")],
        )
        analysis = MinuteAnalysis(items=[
            MeetingItem(
                category="acuerdo",
                description="Punto aprobado.",
                review_status="aprobado",
            ),
            MeetingItem(
                category="informativo",
                description="Punto descartado.",
                review_status="descartado",
            ),
        ])
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "minuta.docx"
            generate_ash_docx(analysis, metadata, output, ROOT / "assets" / "logo_ash.png")
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            text += "\n" + "\n".join(
                cell.text for table in document.tables for row in table.rows for cell in row.cells
            )
            self.assertIn("Punto aprobado", text)
            self.assertNotIn("Punto descartado", text)


if __name__ == "__main__":
    unittest.main()
