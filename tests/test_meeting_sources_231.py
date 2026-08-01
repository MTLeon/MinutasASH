from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from unittest.mock import patch

from docx import Document

from src.meeting_sources import (
    create_text_source,
    infer_source_type,
    parse_text_transcript,
    read_meeting_source,
)


class FlexibleMeetingSources231Tests(unittest.TestCase):
    def test_text_with_speaker_and_timestamp_is_normalized(self):
        segments = parse_text_transcript(
            "[00:01:20] Carlos Pérez: Yo voy a consultar el estado del pago.\n"
            "Ana Gómez: Estamos a la espera de la orden de compra."
        )
        self.assertEqual(len(segments), 2)
        self.assertEqual(segments[0].speaker, "Carlos Pérez")
        self.assertEqual(segments[0].start, "00:01:20.000")
        self.assertIn("estado del pago", segments[0].text)
        self.assertEqual(segments[1].speaker, "Ana Gómez")

    def test_plain_continuation_is_not_lost(self):
        segments = parse_text_transcript(
            "Mauricio: Debemos revisar el cierre técnico.\n"
            "También falta el protocolo firmado."
        )
        self.assertEqual(len(segments), 1)
        self.assertIn("También falta", segments[0].text)

    def test_docx_source_reads_paragraphs_and_tables(self):
        with TemporaryDirectory() as temp:
            path = Path(temp) / "reunion.docx"
            document = Document()
            document.add_paragraph("Mauricio: Se revisará el proyecto 3261.")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Iván"
            table.cell(0, 1).text = "Queda pendiente la orden de compra."
            document.save(path)
            source = read_meeting_source(path)
        self.assertEqual(source.source_type, "docx")
        self.assertGreaterEqual(len(source.segments), 1)
        combined = " ".join(segment.text for segment in source.segments)
        self.assertIn("orden de compra", combined)

    def test_manual_text_is_saved_as_a_local_source(self):
        with TemporaryDirectory() as temp, patch(
            "src.meeting_sources.drafts_dir", return_value=Path(temp)
        ):
            source = create_text_source(
                "Mauricio: Yo enviaré el informe el lunes.",
                source_type="pasted",
                suggested_name="reunión interna",
            )
            self.assertTrue(source.path.is_file())
            self.assertEqual(source.source_type, "pasted")
            self.assertEqual(source.quality, "media")
            self.assertIn("enviaré", source.path.read_text(encoding="utf-8"))

    def test_unknown_extension_is_rejected(self):
        with self.assertRaises(ValueError):
            infer_source_type("reunion.pdf")


if __name__ == "__main__":
    unittest.main()
